#!/usr/bin/env python3
"""Export a repo-format resume markdown file to DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from export_resume_web_pdf import BulletItem, ExperienceEntry, parse_bold_prefix, parse_resume


INK = RGBColor(0x0A, 0x0A, 0x0A)
MUTED = RGBColor(0x4A, 0x4A, 0x4A)
RULE = "D9D9D9"


def rgb_from_hex(value: str) -> RGBColor:
    cleaned = value.strip().lstrip("#")
    if not re.fullmatch(r"[0-9A-Fa-f]{6}", cleaned):
        cleaned = "8c1515"
    return RGBColor(int(cleaned[0:2], 16), int(cleaned[2:4], 16), int(cleaned[4:6], 16))


def set_font(run, size: int, color: RGBColor = INK, bold: bool = False) -> None:
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def add_bottom_border(paragraph, color: str = RULE, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_top_border(paragraph, color: str = RULE, size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color)
    borders.append(top)


def add_inline_markdown(paragraph, text: str, size: int = 10, color: RGBColor = INK) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, size, color, bold=True)
        else:
            run = paragraph.add_run(part.replace("*", ""))
            set_font(run, size, color)


def add_section_heading(doc: Document, title: str, theme: RGBColor) -> None:
    paragraph = doc.add_paragraph()
    add_top_border(paragraph)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(title.upper())
    set_font(run, 12, theme, bold=True)


def add_bullet(doc: Document, bullet: BulletItem, depth: int = 0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375 + depth * 0.2)
    paragraph.paragraph_format.first_line_indent = Inches(-0.1875)
    paragraph.paragraph_format.space_after = Pt(3)
    marker = paragraph.add_run("•  ")
    set_font(marker, 10)
    add_inline_markdown(paragraph, bullet.text, 10)
    for child in bullet.children:
        add_bullet(doc, child, depth + 1)


def add_experience(doc: Document, entry: ExperienceEntry) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(1)
    run = heading.add_run(entry.role)
    set_font(run, 12, bold=True)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(3)
    meta_text = " · ".join(part for part in [entry.company, entry.context, entry.dates] if part)
    run = meta.add_run(meta_text)
    set_font(run, 10, MUTED)

    for bullet in entry.bullets:
        add_bullet(doc, bullet)


def add_blocks(doc: Document, blocks: list[list[str]]) -> None:
    for block in blocks:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        add_inline_markdown(paragraph, " ".join(line.strip() for line in block if line.strip()), 10)


def add_skills(doc: Document, rows: list[str]) -> None:
    for row in rows:
        label, rest = parse_bold_prefix(row)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        if label:
            run = paragraph.add_run(label)
            set_font(run, 10, bold=True)
            add_inline_markdown(paragraph, rest, 10)
        else:
            add_inline_markdown(paragraph, row, 10)


def export_docx(input_md: Path, output_docx: Path, theme_hex: str) -> None:
    data = parse_resume(input_md)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    theme = rgb_from_hex(theme_hex)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(data.name)
    set_font(run, 24, theme, bold=True)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(4)
    run = tagline.add_run(data.tagline)
    set_font(run, 12, MUTED)

    for summary in data.summary:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(6)
        add_inline_markdown(paragraph, summary, 10)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_before = Pt(2)
    contact.paragraph_format.space_after = Pt(8)
    add_top_border(contact)
    add_bottom_border(contact)
    run = contact.add_run(" · ".join(data.contacts))
    set_font(run, 9, MUTED)

    add_section_heading(doc, "Professional Experience", theme)
    for entry in data.experience:
        add_experience(doc, entry)

    add_section_heading(doc, "Core Skills", theme)
    add_skills(doc, data.core_skills)

    add_section_heading(doc, "Education", theme)
    add_blocks(doc, data.education)

    if data.additional:
        add_section_heading(doc, "Additional Experience & Languages", theme)
        add_blocks(doc, data.additional)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)


def main() -> None:
    parser = argparse.ArgumentParser(description="Export resume markdown to DOCX.")
    parser.add_argument("input_md")
    parser.add_argument("output_docx")
    parser.add_argument("--theme", default="#8c1515")
    args = parser.parse_args()
    export_docx(Path(args.input_md), Path(args.output_docx), args.theme)


if __name__ == "__main__":
    main()
